from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for border_name, value in kwargs.items():
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tcBorders.append(border)

        if value is None:
            if border is not None:
                tcBorders.remove(border)
        else:
            border.set(qn('w:val'), value.get('val', 'single'))
            border.set(qn('w:sz'), value.get('sz', '4'))
            border.set(qn('w:space'), value.get('space', '0'))
            border.set(qn('w:color'), value.get('color', 'auto'))


def set_cell_background_color(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)


def fill_table_data(table, data):
    for i, (label, value) in enumerate(data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

def format_table(table):
    for row_idx in range(3):  
        set_cell_background_color(table.cell(row_idx, 0), "D9E1F2")
        set_cell_background_color(table.cell(row_idx, 1), "D9E1F2")

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "000000"},
                bottom={"val": "single", "sz": "4", "color": "000000"},
                left={"val": "single", "sz": "4", "color": "000000"},
                right={"val": "single", "sz": "4", "color": "000000"},
            )

def align_headings(doc):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def create_document():

    doc = Document()

    # Adicionando título e subtítulo
    doc.add_heading('MATRIZ DE PLANEJAMENTO E DESIGN EDUCACIONAL', level=0)
    doc.add_heading('1. DADOS GERAIS', level=1)

    # Dados para preencher na tabela
    dadosGerais = [
        ('Curso', 'CURSO TÉCNICO DE EVENTOS'),
        ('Disciplina', 'ASPECTOS SOCIOCULTURAIS EM EVENTOS'),
        ('Semestre', '2024.2'),
        ('Período letivo de planejamento', '2024.1'),
        ('Período letivo de oferta', '2024.2'),
        ('Formato de oferta da disciplina', 'MODULAR'),
        ('Professor(a):', 'CAROLINA CASTELO BRANCO')
    ]

   
    tabela = doc.add_table(rows=7, cols=2)
    fill_table_data(tabela,dadosGerais)

    format_table(tabela)

    align_headings(doc)

    doc.save('matriz_planejamento_design_educacional.docx')
    print("Documento criado ou atualizado com sucesso!")

create_document()
